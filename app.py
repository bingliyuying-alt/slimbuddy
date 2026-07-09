"""
SlimBuddy - Your Weight Loss Buddy
FastAPI backend + static frontend
"""
import os, sys, json, asyncio, urllib.request
from datetime import datetime, date, timedelta
from contextlib import asynccontextmanager

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import HTMLResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
import uvicorn

from slim.database import (
    init_db, get_latest_weight, save_weight, get_weight_trend,
    get_today_plan, save_today_plan, get_week_schedule, save_week_schedule,
    get_today_checkins, save_checkin, get_last_checkin_date,
    get_week_stats, save_body_state, get_ai_memories, save_ai_memory,
    get_user_profile, get_db,
)
from slim.planner import generate_daily_plan, generate_weekly_prompts, get_week_monday
from slim.scale import ScaleReader
from slim.health import calc_bmi, calc_bmr, calc_tdee, calc_target_calories, bmi_label

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")

@asynccontextmanager
async def lifespan(app: FastAPI):
    init_db()
    yield

app = FastAPI(title="SlimBuddy", version="0.1.0", lifespan=lifespan)

app.mount("/static", StaticFiles(directory=os.path.join(BASE_DIR, "static")), name="static")
try:
    app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")
except:
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")


@app.get("/api/today")
async def api_today():
    today = date.today().strftime("%Y-%m-%d")
    plan = get_today_plan(today)
    weight = get_latest_weight()
    checkins = get_today_checkins(today)
    last_checkin = get_last_checkin_date()
    gap_days = None
    if last_checkin and last_checkin != today:
        last_date = datetime.strptime(last_checkin, "%Y-%m-%d").date()
        gap_days = (date.today() - last_date).days
    return {
        "date": today, "weekday": ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"][date.today().weekday()],
        "weight": weight["weight_kg"] if weight else None,
        "plan": plan, "checkins": checkins,
        "checkin_complete": len([c for c in checkins if c["completed"]]) >= 3,
        "gap_days": gap_days,
    }

@app.get("/api/weight/trend")
async def api_weight_trend(days: int = 30):
    return {"data": get_weight_trend(days)}

@app.post("/api/weight")
async def api_save_weight(weight: float = Form(...), notes: str = Form("")):
    save_weight(weight, source="manual", notes=notes)
    conn = get_db()
    conn.execute("DELETE FROM daily_plans WHERE plan_date=date('now','localtime')")
    conn.commit(); conn.close()
    return {"ok": True}

@app.get("/api/scale/scan")
async def api_scale_scan():
    try:
        reader = ScaleReader()
        weight = await reader.wait_for_reading(timeout=30)
        if weight:
            save_weight(weight, source="scale")
            # 体重更新后清除计划缓存，下次打开基于新体重生成
            conn = get_db()
            conn.execute("DELETE FROM daily_plans WHERE plan_date=date('now','localtime')")
            conn.commit(); conn.close()
            return {"ok": True, "weight": weight}
        return {"ok": False, "message": "No scale found"}
    except Exception as e:
        return {"ok": False, "message": str(e)}

@app.get("/api/plan/today")
async def api_get_plan(weather: str = None):
    today = date.today().strftime("%Y-%m-%d")
    plan = get_today_plan(today)
    now = datetime.now()
    # 凌晨4点后自动翻篇：如果缓存是昨天的，重新生成
    if plan is not None and now.hour >= 4:
        plan_date = plan["plan_date"]
        if plan_date < today:
            plan = None
    if plan is None:
        latest = get_latest_weight()
        w = latest["weight_kg"] if latest else None
        g = generate_daily_plan(today, weather=weather, weight_kg=w)
        save_today_plan(today,
            json.dumps(g["breakfast"],ensure_ascii=False),
            json.dumps(g["lunch"],ensure_ascii=False),
            json.dumps(g["dinner"],ensure_ascii=False),
            json.dumps(g["exercise"],ensure_ascii=False),
            json.dumps({"motto":g["motto"],"schedule":g["schedule"]},ensure_ascii=False))
        return g
    return {
        "date": today, "weekday": ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"][date.today().weekday()],
        "weight": get_latest_weight(),
        "breakfast": json.loads(plan["breakfast"]) if plan["breakfast"] else None,
        "lunch": json.loads(plan["lunch"]) if plan["lunch"] else None,
        "dinner": json.loads(plan["dinner"]) if plan["dinner"] else None,
        "exercise": json.loads(plan["exercise"]) if plan["exercise"] else None,
        "extra": json.loads(plan["extra_notes"]) if plan["extra_notes"] else {},
        "modified": bool(plan["modified"]),
    }

@app.post("/api/plan/update")
async def api_update_plan(meal: str = Form(...), name: str = Form(""), detail: str = Form(""),
                          calories: int = Form(None)):
    today = date.today().strftime("%Y-%m-%d")
    plan = get_today_plan(today)
    if not plan: return {"ok": False}
    bf = json.loads(plan["breakfast"]) if plan["breakfast"] else {}
    lu = json.loads(plan["lunch"]) if plan["lunch"] else {}
    di = json.loads(plan["dinner"]) if plan["dinner"] else {}
    ex = json.loads(plan["exercise"]) if plan["exercise"] else {}
    target = {"breakfast":bf,"lunch":lu,"dinner":di,"exercise":ex}.get(meal,{})
    if name: target["name"] = name
    if detail: target["detail"] = detail
    if calories is not None: target["calories"] = calories
    save_today_plan(today, json.dumps(bf,ensure_ascii=False), json.dumps(lu,ensure_ascii=False),
                    json.dumps(di,ensure_ascii=False), json.dumps(ex,ensure_ascii=False), plan["extra_notes"])
    from slim.database import get_db
    conn = get_db(); conn.execute("UPDATE daily_plans SET modified=1 WHERE plan_date=?",(today,)); conn.commit(); conn.close()
    return {"ok": True}

@app.post("/api/checkin")
async def api_checkin(meal_type: str = Form(...), note: str = Form(""), 
                      photos: list[UploadFile] = File(default=None)):
    today = date.today().strftime("%Y-%m-%d")
    # 去重：同日期同餐型只保留最新一条
    conn = get_db()
    conn.execute("DELETE FROM checkins WHERE checkin_date=? AND meal_type=?", (today, meal_type))
    conn.commit(); conn.close()
    paths = []
    MAX_PHOTOS = 9
    if photos:
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        for i, photo in enumerate(photos[:MAX_PHOTOS]):
            if photo and photo.filename:
                fn = f"{today}_{meal_type}_{ts}_{i}.jpg"
                pp = os.path.join("uploads", fn)
                with open(os.path.join(BASE_DIR, pp), "wb") as f:
                    f.write(await photo.read())
                paths.append(pp)
    photo_path = json.dumps(paths) if paths else None
    save_checkin(today, meal_type, photo_path, note, completed=1)
    if note: save_ai_memory(f"[{meal_type}] {note}", "insight")
    return {"ok": True, "photos": len(paths)}

@app.get("/api/week/schedule")
async def api_get_week_schedule():
    ws = get_week_monday()
    data = get_week_schedule(ws)
    prompts = generate_weekly_prompts()
    sm = {s["schedule_date"]: s.get("plan_notes","") for s in data}
    for p in prompts: p["notes"] = sm.get(p["date"], "")
    return {"week_start": ws, "days": prompts}

@app.post("/api/week/schedule")
async def api_save_week_schedule(request: Request):
    body = await request.json()
    ws = body.get("week_start", get_week_monday())
    for d in body.get("days", []):
        save_week_schedule(ws, d["day_index"], d["date"], d.get("notes", ""))
    return {"ok": True}

@app.get("/api/week/stats")
async def api_week_stats():
    ws = get_week_monday()
    stats = get_week_stats(ws)
    mems = get_ai_memories(7)
    profile = get_user_profile()
    return {"week_start": ws, "stats": stats, "memories": [m["content"] for m in mems],
            "profile_weight": profile["start_weight_kg"] if profile else None}

@app.post("/api/body/state")
async def api_save_body_state(sleep_hours: float = Form(None), sleep_quality: str = Form(None),
                               mood: str = Form(None), notes: str = Form(None)):
    save_body_state(date.today().strftime("%Y-%m-%d"), sleep_hours, sleep_quality, mood, notes)
    return {"ok": True}

@app.get("/api/gap/recover")
async def api_gap_recover():
    last = get_last_checkin_date()
    today = date.today().strftime("%Y-%m-%d")
    if last and last != today:
        gap = (date.today() - datetime.strptime(last, "%Y-%m-%d").date()).days
        return {"gap": gap, "messages": [f"Hey! {gap} days since I last saw you. No worries at all."]}
    return {"gap": 0, "messages": []}

@app.get("/api/weather")
async def api_weather():
    try:
        req = urllib.request.Request("https://wttr.in/Wuxi?format=%C+%t&lang=zh",
                                      headers={"User-Agent": "curl/8.0"})
        with urllib.request.urlopen(req, timeout=5) as resp:
            raw = resp.read().decode("utf-8").strip()
            import re
            if raw.startswith("<"):
                m = re.search(r'<div class="term-container">(.+?)</div>', raw, re.DOTALL)
                raw = m.group(1).strip() if m else re.sub(r'<[^>]+>', '', raw).strip()
            return {"weather": raw}
    except:
        return {"weather": "--"}


@app.get("/api/export")
async def api_export(days: int = 7):
    """导出 Word 日报 (.docx)"""
    import io
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    conn = get_db()
    end_date = date.today().strftime("%Y-%m-%d")
    start_date = (date.today() - timedelta(days=days)).strftime("%Y-%m-%d")

    # 获取打卡记录（按日期排序）
    checkins = conn.execute(
        "SELECT checkin_date, meal_type, note, photo_path FROM checkins "
        "WHERE checkin_date >= ? AND checkin_date <= ? "
        "ORDER BY checkin_date ASC, created_at ASC",
        (start_date, end_date)
    ).fetchall()

    # 获取体重记录 (recorded_at 含时间，用 date() 比对)
    weights = conn.execute(
        "SELECT recorded_at, weight_kg FROM weight_records "
        "WHERE date(recorded_at) >= ? AND date(recorded_at) <= ? ORDER BY recorded_at ASC",
        (start_date, end_date)
    ).fetchall()

    # 获取身体状态（心情）
    body_states = conn.execute(
        "SELECT recorded_date, sleep_hours, sleep_quality, mood, notes FROM body_state "
        "WHERE recorded_date >= ? AND recorded_date <= ? ORDER BY recorded_date ASC",
        (start_date, end_date)
    ).fetchall()

    # 获取每日计划
    plans = conn.execute(
        "SELECT plan_date, breakfast, lunch, dinner, exercise FROM daily_plans "
        "WHERE plan_date >= ? AND plan_date <= ? ORDER BY plan_date ASC",
        (start_date, end_date)
    ).fetchall()

    # 获取用户信息
    profile = conn.execute("SELECT * FROM user_profile WHERE id=1").fetchone()
    conn.close()

    # 按日期分组打卡
    from collections import defaultdict
    ck_by_date = defaultdict(list)
    for row in checkins:
        ck_by_date[row["checkin_date"]].append(dict(row))

    ws_by_date = {}
    for row in weights:
        d = row["recorded_at"][:10]
        ws_by_date[d] = row["weight_kg"]

    bs_by_date = {}
    for row in body_states:
        bs_by_date[row["recorded_date"]] = dict(row)

    pl_by_date = {}
    for row in plans:
        pl_by_date[row["plan_date"]] = dict(row)

    # 生成文档
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    # 标题
    title = doc.add_heading("SlimBuddy 减脂日报", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"日期范围：{start_date} 至 {end_date}")
    if profile:
        info = f"身高：{profile['height_cm']}cm | 年龄：{profile['age']}岁"
        doc.add_paragraph(info)
    doc.add_paragraph("")

    # 体重总览
    doc.add_heading("体重趋势", level=2)
    if weights:
        for w in weights:
            doc.add_paragraph(f"{w['recorded_at'][:10]}  ➜  {w['weight_kg']} kg")
    else:
        doc.add_paragraph("本周暂无体重记录")
    doc.add_paragraph("")

    # 每日详情
    doc.add_heading("每日打卡详情", level=2)

    all_dates = sorted(set(list(ck_by_date.keys()) + list(ws_by_date.keys()) + list(bs_by_date.keys())))
    if not all_dates:
        # 用日期范围生成
        d0 = datetime.strptime(start_date, "%Y-%m-%d").date()
        for i in range(days):
            all_dates.append((d0 + timedelta(days=i)).strftime("%Y-%m-%d"))

    meal_labels = {"breakfast": "早餐", "lunch": "午餐", "dinner": "晚餐", "exercise": "运动"}
    meal_emojis = {"breakfast": "🌅", "lunch": "☀️", "dinner": "🌙", "exercise": "🏃"}

    for d in all_dates:
        cks = ck_by_date.get(d, [])
        w = ws_by_date.get(d)
        b = bs_by_date.get(d)
        pl = pl_by_date.get(d)

        date_label = f"📅 {d}"
        if pl:
            weekday = ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"][datetime.strptime(d,"%Y-%m-%d").weekday()]
            date_label += f" ({weekday})"
        doc.add_heading(date_label, level=3)

        if w:
            doc.add_paragraph(f"⚖️ 体重：{w} kg")

        # 计划 / 打卡
        for meal_type in ["breakfast", "lunch", "dinner", "exercise"]:
            lines = []
            # 从计划中获取名称和热量
            if pl:
                plan_json = pl.get(meal_type)
                if plan_json:
                    try:
                        plan_data = json.loads(plan_json)
                        name = plan_data.get("name", "")
                        detail = plan_data.get("detail", "")
                        cal = plan_data.get("calories", "")
                        cal_burn = plan_data.get("calories_burn", "")
                        if name:
                            line = f"{meal_emojis[meal_type]} {name}"
                            if cal:
                                line += f" — {cal} kcal"
                            if cal_burn:
                                line += f" — 消耗 {cal_burn} kcal"
                            if detail and detail != name:
                                line += f"  ({detail})"
                            lines.append(line)
                    except:
                        pass

            # 打卡备注
            day_cks = [c for c in cks if c["meal_type"] == meal_type]
            for ck in day_cks:
                if ck["note"]:
                    lines.append(f"   备注：{ck['note']}")
                if ck["photo_path"]:
                    try:
                        pics = json.loads(ck["photo_path"])
                        lines.append(f"   📷 {len(pics)} 张照片")
                    except:
                        lines.append("   📷 有照片")

            for line in lines:
                doc.add_paragraph(line, style="List Bullet")

            if day_cks:
                doc.add_paragraph("✅ 已打卡", style="List Bullet")
            else:
                doc.add_paragraph(f"{meal_emojis[meal_type]} {meal_labels[meal_type]}：未打卡", style="List Bullet")

        # 心情
        if b:
            mood_text = f"💭 心情：{b.get('mood','未记录')}"
            if b.get("sleep_hours"):
                mood_text += f"  |  睡眠：{b['sleep_hours']}h ({b.get('sleep_quality','')})"
            if b.get("notes"):
                mood_text += f"  |  {b['notes']}"
            doc.add_paragraph(mood_text)

        doc.add_paragraph("")

    # AI 洞察
    mems = get_ai_memories(days)
    if mems:
        doc.add_heading("AI 洞察", level=2)
        for m in mems:
            doc.add_paragraph(m["content"], style="List Bullet")

    doc.add_paragraph("")
    p_footer = doc.add_paragraph("—— SlimBuddy 减脂搭子 · 温柔陪伴，不评判 ——")
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 输出为字节流
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"SlimBuddy_{start_date}_{end_date}.docx"
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/api/history")
async def api_history(days: int = 7):
    conn = get_db()
    rows = conn.execute(
        "SELECT checkin_date, meal_type, note, photo_path, completed FROM checkins "
        "WHERE checkin_date >= date('now','localtime',?) "
        "ORDER BY checkin_date DESC, created_at DESC LIMIT 50",
        (f'-{days} days',)
    ).fetchall()
    conn.close()
    return {"history": [dict(r) for r in rows]}

@app.get("/api/settings")
async def api_get_settings():
    profile = get_user_profile()
    if not profile:
        return {"ok": False}
    weight = get_latest_weight()
    w = weight["weight_kg"] if weight else None
    bmi = calc_bmi(w, profile.get("height_cm"))
    bmr = calc_bmr(w, profile.get("height_cm"), profile.get("age"))
    tdee = calc_tdee(bmr, profile.get("activity_level","light")) if bmr else None
    target = calc_target_calories(tdee) if tdee else None
    return {
        "ok": True,
        "profile": profile,
        "current_weight": w,
        "bmi": bmi, "bmi_label": bmi_label(bmi),
        "bmr": bmr, "tdee": tdee, "target_calories": target
    }

@app.post("/api/settings")
async def api_save_settings(request: Request):
    body = await request.json()
    conn = get_db()
    fields = ["height_cm","age","allergies","conditions","medications","activity_level"]
    for f in fields:
        if f in body and body[f] is not None:
            conn.execute(f"UPDATE user_profile SET {f}=? WHERE id=1", (str(body[f]),))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.get("/", response_class=HTMLResponse)
async def index():
    hp = os.path.join(BASE_DIR, "static", "index.html")
    if os.path.exists(hp):
        with open(hp, "r", encoding="utf-8") as f: content = f.read()
        return Response(content=content, media_type="text/html", headers={"Cache-Control": "no-cache"})
    return HTMLResponse("<h1>SlimBuddy</h1>")

if __name__ == "__main__":
    print("\n    ================================")
    print("       SlimBuddy  http://localhost:8765")
    print("    ================================\n")
    uvicorn.run(app, host="0.0.0.0", port=8765, log_level="info")
